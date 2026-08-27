import math
import base64
import json
import os
import re
import sqlite3
import csv
import urllib.error
import urllib.request
import tempfile
from urllib.parse import urlparse
from collections import Counter
from datetime import datetime, timezone
from functools import wraps
from io import BytesIO

from flask import Flask, Response, flash, jsonify, redirect, render_template, request, send_file, session, url_for
from pypdf import PdfReader
from docx import Document
from werkzeug.exceptions import RequestEntityTooLarge
from werkzeug.security import check_password_hash, generate_password_hash
from werkzeug.utils import secure_filename

try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass


app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 10 * 1024 * 1024
app.config["SECRET_KEY"] = os.environ.get("VERITASCHECK_SECRET", "dev-only-change-this-key")
IS_VERCEL = bool(os.environ.get("VERCEL"))
RUNTIME_DATA_ROOT = os.environ.get("VERITASCHECK_DATA_ROOT") or (tempfile.gettempdir() if IS_VERCEL else app.root_path)
DATABASE_PATH = os.environ.get("VERITASCHECK_DATABASE_PATH") or os.path.join(RUNTIME_DATA_ROOT, "database", "veritascheck.db")
UPLOADS_ROOT = os.environ.get("VERITASCHECK_UPLOADS_PATH") or os.path.join(RUNTIME_DATA_ROOT, "uploads")

ALLOWED_EXTENSIONS = {"txt", "pdf", "docx", "png", "jpg", "jpeg", "webp"}
IMAGE_EXTENSIONS = {"png", "jpg", "jpeg", "webp"}
MIN_WORDS = 20
MATCH_THRESHOLD = 0.55
MAX_MATCHES = 20
OPENROUTER_API_KEY = os.environ.get("OPENROUTER_API_KEY", "")
OPENROUTER_MODEL = os.environ.get("OPENROUTER_MODEL", "openrouter/auto")
OPENROUTER_VISION_MODEL = os.environ.get("OPENROUTER_VISION_MODEL", "openrouter/free")
MODEL_PATH = os.environ.get("VERITASCHECK_MODEL_PATH", os.path.join(app.root_path, "veritas_trained_model"))
SEMANTIC_MODEL = None
SEMANTIC_UTIL = None
if os.path.isdir(MODEL_PATH):
    try:
        from sentence_transformers import SentenceTransformer, util as sentence_util
        SEMANTIC_MODEL = SentenceTransformer(MODEL_PATH)
        SEMANTIC_UTIL = sentence_util
    except (ImportError, OSError, ValueError) as model_error:
        app.logger.warning("Semantic model could not be loaded; using cosine baseline: %s", model_error)


@app.after_request
def prevent_stale_development_assets(response):
    """Keep rapidly changing project UI from being hidden by browser caches."""
    response.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, max-age=0"
    response.headers["Pragma"] = "no-cache"
    response.headers["Expires"] = "0"
    return response


class ClosingConnection(sqlite3.Connection):
    def __exit__(self, exc_type, exc_value, traceback):
        try:
            return super().__exit__(exc_type, exc_value, traceback)
        finally:
            self.close()


def get_db():
    os.makedirs(os.path.dirname(DATABASE_PATH), exist_ok=True)
    connection = sqlite3.connect(DATABASE_PATH, factory=ClosingConnection)
    connection.row_factory = sqlite3.Row
    return connection


def init_db():
    with get_db() as database:
        database.executescript("""
            CREATE TABLE IF NOT EXISTS users (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                name TEXT NOT NULL,
                email TEXT NOT NULL UNIQUE,
                password_hash TEXT NOT NULL,
                created_at TEXT NOT NULL
            );
            CREATE TABLE IF NOT EXISTS analyses (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER NOT NULL,
                filename TEXT NOT NULL,
                words INTEGER NOT NULL,
                pages INTEGER NOT NULL,
                similarity REAL NOT NULL,
                matches INTEGER NOT NULL,
                result_json TEXT NOT NULL,
                created_at TEXT NOT NULL,
                file_path TEXT,
                ai_score REAL,
                grammar_score REAL,
                FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE CASCADE
            );
        """)
        columns = {row["name"] for row in database.execute("PRAGMA table_info(analyses)").fetchall()}
        if "file_path" not in columns:
            database.execute("ALTER TABLE analyses ADD COLUMN file_path TEXT")
        if "ai_score" not in columns:
            database.execute("ALTER TABLE analyses ADD COLUMN ai_score REAL")
        if "grammar_score" not in columns:
            database.execute("ALTER TABLE analyses ADD COLUMN grammar_score REAL")


def login_required(view):
    @wraps(view)
    def wrapped(*args, **kwargs):
        if not session.get("user_id"):
            flash("Please sign in to access your saved reports.", "info")
            return redirect(url_for("login"))
        return view(*args, **kwargs)
    return wrapped


@app.context_processor
def current_user_context():
    if not session.get("user_id"):
        return {"current_user": None}
    with get_db() as database:
        user = database.execute("SELECT id, name, email FROM users WHERE id = ?", (session["user_id"],)).fetchone()
    return {"current_user": user}


def tokenize(text):
    return re.findall(r"[\w']+", text.lower(), flags=re.UNICODE)


def split_sentences(text):
    return [part.strip() for part in re.split(r"(?<=[.!?])\s+|[\r\n]+", text) if part.strip()]


def locate_text(full_text, sentence):
    position = full_text.find(sentence)
    return full_text[:max(position, 0)].count("\n") + 1


def cosine_similarity(left, right):
    """Transparent baseline; replace vectors here with embedding vectors later."""
    left_counts, right_counts = Counter(tokenize(left)), Counter(tokenize(right))
    if not left_counts or not right_counts:
        return 0.0
    dot = sum(value * right_counts.get(term, 0) for term, value in left_counts.items())
    left_mag = math.sqrt(sum(value * value for value in left_counts.values()))
    right_mag = math.sqrt(sum(value * value for value in right_counts.values()))
    return dot / (left_mag * right_mag) if left_mag and right_mag else 0.0


def calculate_similarity(document, reference):
    if not reference.strip():
        return 0.0
    if SEMANTIC_MODEL is not None:
        embeddings = SEMANTIC_MODEL.encode([document, reference], convert_to_tensor=True)
        return round(max(0.0, min(1.0, SEMANTIC_UTIL.cos_sim(embeddings[0], embeddings[1]).item())) * 100, 1)
    return round(cosine_similarity(document, reference) * 100, 1)


def ensure_source_evidence(document_pages, sources, matches):
    """Keep a strongest, explicitly scored passage for each discovered source."""
    covered = {match["source_name"] for match in matches}
    for source in sources:
        if source["name"] in covered:
            continue
        best = None
        for document_page, document_text in enumerate(document_pages, 1):
            for sentence in [item for item in split_sentences(document_text) if len(tokenize(item)) >= 5]:
                for source_page, source_text in enumerate(source["pages"], 1):
                    for candidate in [item for item in split_sentences(source_text) if len(tokenize(item)) >= 5]:
                        score = cosine_similarity(sentence, candidate)
                        if best is None or score > best[0]:
                            best = (score, sentence, candidate, document_page, locate_text(document_text, sentence),
                                    source_page, locate_text(source_text, candidate))
        # Below 35% is too weak to highlight as useful evidence.
        if best and best[0] >= 0.35:
            score, sentence, candidate, document_page, document_line, source_page, source_line = best
            matches.append({"text": sentence, "score": round(score * 100, 1), "source": candidate,
                            "source_name": source["name"], "source_url": source.get("url", ""),
                            "source_page": source_page, "source_line": source_line,
                            "document_page": document_page, "document_line": document_line,
                            "evidence_type": "closest_verified_passage"})
    return matches


def find_matches(document_pages, sources):
    if not sources:
        return []
    matches = []
    source_candidates = []
    for source in sources:
        for page_number, page_text in enumerate(source["pages"], 1):
            for candidate in [s for s in split_sentences(page_text) if len(tokenize(s)) >= 5]:
                source_candidates.append({"text": candidate, "name": source["name"], "url": source.get("url", ""), "page": page_number,
                                          "line": locate_text(page_text, candidate)})
    if SEMANTIC_MODEL is not None and source_candidates:
        document_candidates = []
        for page_number, page_text in enumerate(document_pages, 1):
            for sentence in [s for s in split_sentences(page_text) if len(tokenize(s)) >= 5]:
                document_candidates.append({"text": sentence, "page": page_number,
                                            "line": locate_text(page_text, sentence)})
        if document_candidates:
            left_embeddings = SEMANTIC_MODEL.encode([item["text"] for item in document_candidates], convert_to_tensor=True)
            right_embeddings = SEMANTIC_MODEL.encode([item["text"] for item in source_candidates], convert_to_tensor=True)
            matrix = SEMANTIC_UTIL.cos_sim(left_embeddings, right_embeddings)
            for row_index, document_item in enumerate(document_candidates):
                best_score, best_index = matrix[row_index].max(dim=0)
                score = float(best_score.item())
                if score >= MATCH_THRESHOLD:
                    source_item = source_candidates[int(best_index.item())]
                    matches.append({"text": document_item["text"], "score": round(score * 100, 1),
                                    "source": source_item["text"], "source_name": source_item["name"],
                                    "source_url": source_item["url"],
                                    "source_page": source_item["page"], "source_line": source_item["line"],
                                    "document_page": document_item["page"], "document_line": document_item["line"]})
            matches = ensure_source_evidence(document_pages, sources, matches)
            return sorted(matches, key=lambda item: item["score"], reverse=True)[:max(MAX_MATCHES, len(sources))]
    for document_page, document_text in enumerate(document_pages, 1):
        for sentence in [s for s in split_sentences(document_text) if len(tokenize(s)) >= 5]:
            best_source, best_score, best_name, best_url, best_page, best_line = "", 0.0, "", "", None, 1
            for source in sources:
                for page_number, page_text in enumerate(source["pages"], 1):
                    for candidate in [s for s in split_sentences(page_text) if len(tokenize(s)) >= 5]:
                        score = cosine_similarity(sentence, candidate)
                        if score > best_score:
                            best_source, best_score = candidate, score
                            best_name, best_url, best_page = source["name"], source.get("url", ""), page_number
                            best_line = locate_text(page_text, candidate)
            if best_score >= MATCH_THRESHOLD:
                matches.append({"text": sentence, "score": round(best_score * 100, 1), "source": best_source,
                                "source_name": best_name, "source_url": best_url, "source_page": best_page, "source_line": best_line,
                                "document_page": document_page, "document_line": locate_text(document_text, sentence)})
    matches = ensure_source_evidence(document_pages, sources, matches)
    return sorted(matches, key=lambda item: item["score"], reverse=True)[:max(MAX_MATCHES, len(sources))]


def extract_image_text(data, extension):
    if not OPENROUTER_API_KEY:
        raise ValueError("Image OCR is not configured. Add OPENROUTER_API_KEY and redeploy the application.")
    mime = "image/jpeg" if extension in {"jpg", "jpeg"} else f"image/{extension}"
    encoded = base64.b64encode(data).decode("ascii")
    payload = {"model": OPENROUTER_VISION_MODEL, "temperature": 0, "max_tokens": 6000,
               "messages": [{"role": "user", "content": [
        {"type": "text", "text": (
            "Act as a precise OCR engine, not a summarizer. Transcribe ALL readable text in this document image, "
            "including headings, captions, labels, numbered steps, lists, tables, and small print. Scan from top-left "
            "to bottom-right in natural reading order. Preserve paragraphs and line breaks. Do not describe the image, "
            "omit repeated-looking sections, or shorten the content. Return only the verbatim extracted text without commentary or Markdown fences.")},
        {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{encoded}"}},
    ]}]}
    api_request = urllib.request.Request(
        "https://openrouter.ai/api/v1/chat/completions", data=json.dumps(payload).encode("utf-8"),
        headers={"Authorization": f"Bearer {OPENROUTER_API_KEY}", "Content-Type": "application/json",
                 "HTTP-Referer": "http://127.0.0.1:5000", "X-OpenRouter-Title": "SourceTrace AI"}, method="POST")
    try:
        with urllib.request.urlopen(api_request, timeout=90) as response:
            result = json.loads(response.read().decode("utf-8"))
        content = result["choices"][0]["message"]["content"]
        if isinstance(content, list):
            content = "\n".join(item.get("text", "") for item in content if isinstance(item, dict))
        text = str(content or "").strip()
        if not text:
            raise ValueError("No readable text was detected in this image.")
        return text
    except urllib.error.HTTPError as error:
        try:
            error_data = json.loads(error.read().decode("utf-8"))
            provider_message = str(error_data.get("error", {}).get("message", ""))
        except (json.JSONDecodeError, UnicodeDecodeError, AttributeError):
            provider_message = ""
        if error.code == 401:
            message = "Image OCR authorization failed. Update OPENROUTER_API_KEY in Vercel and redeploy."
        elif error.code == 402:
            message = "Image OCR requires an available OpenRouter model. Enable free-model access or add provider credits."
        elif error.code == 413:
            message = "This image is too large for OCR. Upload a compressed image under 4 MB."
        elif error.code == 429:
            message = "The free image OCR limit is temporarily busy. Wait a minute and try again."
        elif "model" in provider_message.lower() or "image" in provider_message.lower():
            message = "The configured OCR model could not process this image. Set OPENROUTER_VISION_MODEL=openrouter/free and redeploy."
        else:
            message = f"Image OCR provider returned error {error.code}. Please try again shortly."
        app.logger.warning("OpenRouter image OCR failed (%s): %s", error.code, provider_message[:300])
        raise ValueError(message) from error
    except urllib.error.URLError as error:
        app.logger.warning("OpenRouter image OCR network failure: %s", error.reason)
        raise ValueError("Image OCR service could not be reached. Please try again shortly.") from error
    except (KeyError, json.JSONDecodeError, TypeError) as error:
        raise ValueError("Image OCR returned an unreadable response. Please try another image.") from error


def extract_upload(upload):
    filename = secure_filename(upload.filename or "")
    if "." not in filename:
        raise ValueError("Please choose a TXT, PDF, DOCX, or image file.")
    extension = filename.rsplit(".", 1)[1].lower()
    if extension not in ALLOWED_EXTENSIONS:
        raise ValueError("Unsupported file type. Upload TXT, PDF, DOCX, PNG, JPG, JPEG, or WEBP.")
    data = upload.read()
    if not data:
        raise ValueError("The selected file is empty.")
    try:
        if extension == "txt":
            pages = [data.decode("utf-8-sig")]
        elif extension == "docx":
            document = Document(BytesIO(data))
            content = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
            for table in document.tables:
                for row in table.rows:
                    content.append(" | ".join(cell.text.strip() for cell in row.cells))
            pages = ["\n".join(content)]
        elif extension in IMAGE_EXTENSIONS:
            pages = [extract_image_text(data, extension)]
        else:
            pages = []
            for page in PdfReader(BytesIO(data)).pages:
                try:
                    # Layout mode keeps multi-column academic PDFs substantially cleaner.
                    pages.append(page.extract_text(extraction_mode="layout") or "")
                except (TypeError, ValueError):
                    pages.append(page.extract_text() or "")
        return {"name": filename, "pages": pages, "text": "\n".join(pages), "data": data, "extension": extension}
    except ValueError:
        raise
    except Exception as error:
        raise ValueError("We could not read this file. Check that it is valid and not encrypted.") from error


def status_for(score):
    if score < 15:
        return "Low similarity", "success"
    if score < 40:
        return "Moderate similarity", "warning"
    return "High similarity", "danger"


def discover_web_sources(document):
    if not OPENROUTER_API_KEY:
        raise ValueError("Web source discovery is not configured. Add OPENROUTER_API_KEY to .env and restart Flask.")
    candidates = [sentence for sentence in split_sentences(document) if 8 <= len(tokenize(sentence)) <= 55]
    # Search a bounded, representative set to control latency and search cost.
    candidates = sorted(candidates, key=lambda value: len(tokenize(value)), reverse=True)[:12]
    if not candidates:
        raise ValueError("The document does not contain enough searchable sentences.")
    excerpts = "\n".join(f"{index}. {sentence}" for index, sentence in enumerate(candidates, 1))
    payload = {
        "model": OPENROUTER_MODEL,
        "messages": [{"role": "user", "content": (
            "Search the public web for pages containing exact or near-verbatim versions of these document passages. "
            "Use web search repeatedly when useful. Return a brief source list with citations; never invent URLs.\n\n" + excerpts)}],
        "tools": [{"type": "openrouter:web_search", "parameters": {"engine": "auto", "max_total_results": 12}}],
        "temperature": 0,
        "max_tokens": 1200,
    }
    api_request = urllib.request.Request("https://openrouter.ai/api/v1/chat/completions",
        data=json.dumps(payload).encode("utf-8"), method="POST",
        headers={"Authorization": f"Bearer {OPENROUTER_API_KEY}", "Content-Type": "application/json",
                 "HTTP-Referer": "http://127.0.0.1:5000", "X-OpenRouter-Title": "SourceTrace AI"})
    try:
        with urllib.request.urlopen(api_request, timeout=90) as response:
            response_data = json.loads(response.read().decode("utf-8"))
    except urllib.error.HTTPError as error:
        try:
            error_data = json.loads(error.read().decode("utf-8", errors="replace"))
            provider_message = str(error_data.get("error", {}).get("message", ""))
        except (json.JSONDecodeError, AttributeError):
            provider_message = ""
        app.logger.warning("OpenRouter web search failed (%s): %s", error.code, provider_message[:300])
        if error.code == 401:
            message = "Web search authorization failed. Update OPENROUTER_API_KEY in Vercel and redeploy."
        elif error.code == 402:
            message = "The web-search provider has insufficient credits for this request. Add a small OpenRouter balance or choose a lower-cost model."
        elif error.code == 429:
            message = "Web search is temporarily rate-limited. Wait a minute and try again."
        else:
            message = f"Web search provider returned error {error.code}. Please try again shortly."
        raise ValueError(message) from error
    except (urllib.error.URLError, TimeoutError, json.JSONDecodeError) as error:
        raise ValueError("Web source discovery could not connect or returned an invalid response.") from error
    choices = response_data.get("choices") or []
    message = choices[0].get("message", {}) if choices else {}
    annotations = message.get("annotations") or []
    sources, seen = [], set()
    for annotation in annotations:
        citation = annotation.get("url_citation", {}) if annotation.get("type") == "url_citation" else {}
        url, content = citation.get("url", ""), citation.get("content", "")
        if not url or not content or url in seen:
            continue
        seen.add(url)
        domain = urlparse(url).netloc.removeprefix("www.")
        sources.append({"name": citation.get("title") or domain, "domain": domain, "url": url,
                        "pages": [content], "text": content})
    if not sources:
        raise ValueError("Web search completed but returned no verifiable source excerpts for comparison.")
    return sources


def analyze_writing_signals(document):
    """Return non-definitive AI-style and grammar indicators from a hosted model."""
    if not OPENROUTER_API_KEY:
        return {"ai_score": None, "grammar_score": None, "note": "Writing analysis is not configured."}
    sample = document[:16000]
    payload = {"model": OPENROUTER_MODEL, "temperature": 0, "max_tokens": 300,
               "messages": [{"role": "system", "content": (
                   "You assess writing signals. Return only strict JSON with ai_score, grammar_score, and note. "
                   "Both scores are numbers from 0 to 100. ai_score is an experimental stylistic signal, never a claim of authorship. "
                   "grammar_score estimates grammatical correctness and readability. Keep note under 25 words.")},
                   {"role": "user", "content": sample}]}
    api_request = urllib.request.Request("https://openrouter.ai/api/v1/chat/completions",
        data=json.dumps(payload).encode("utf-8"), method="POST",
        headers={"Authorization": f"Bearer {OPENROUTER_API_KEY}", "Content-Type": "application/json",
                 "HTTP-Referer": "http://127.0.0.1:5000", "X-OpenRouter-Title": "SourceTrace AI"})
    try:
        with urllib.request.urlopen(api_request, timeout=90) as response:
            data = json.loads(response.read().decode("utf-8"))
        content = data["choices"][0]["message"]["content"].strip()
        if content.startswith("```"):
            content = re.sub(r"^```(?:json)?\s*|\s*```$", "", content, flags=re.IGNORECASE)
        result = json.loads(content)
        ai_score = round(max(0, min(100, float(result["ai_score"]))), 1)
        grammar_score = round(max(0, min(100, float(result["grammar_score"]))), 1)
        return {"ai_score": ai_score, "grammar_score": grammar_score,
                "note": str(result.get("note", "Experimental model-assisted indicators."))[:240]}
    except (KeyError, IndexError, TypeError, ValueError, json.JSONDecodeError, urllib.error.URLError, TimeoutError):
        return {"ai_score": None, "grammar_score": None, "note": "Writing-signal analysis was unavailable."}


@app.get("/")
def home():
    return render_template("index.html")


@app.get("/checker")
def checker():
    return render_template("checker.html")


@app.get("/about")
def about():
    return render_template("about.html")


@app.route("/register", methods=["GET", "POST"])
def register():
    if request.method == "POST":
        name = request.form.get("name", "").strip()
        email = request.form.get("email", "").strip().lower()
        password = request.form.get("password", "")
        if len(name) < 2 or "@" not in email or len(password) < 8:
            flash("Enter a valid name, email, and password of at least 8 characters.", "danger")
        else:
            try:
                with get_db() as database:
                    cursor = database.execute("INSERT INTO users (name, email, password_hash, created_at) VALUES (?, ?, ?, ?)",
                                              (name, email, generate_password_hash(password), datetime.now(timezone.utc).isoformat()))
                session.clear(); session["user_id"] = cursor.lastrowid
                return redirect(url_for("dashboard"))
            except sqlite3.IntegrityError:
                flash("An account with this email already exists.", "danger")
    return render_template("register.html")


@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        email = request.form.get("email", "").strip().lower()
        with get_db() as database:
            user = database.execute("SELECT * FROM users WHERE email = ?", (email,)).fetchone()
        if user and check_password_hash(user["password_hash"], request.form.get("password", "")):
            session.clear(); session["user_id"] = user["id"]
            return redirect(url_for("dashboard"))
        flash("Email or password is incorrect.", "danger")
    return render_template("login.html")


@app.post("/logout")
def logout():
    session.clear()
    return redirect(url_for("home"))


@app.get("/dashboard")
@login_required
def dashboard():
    with get_db() as database:
        reports = database.execute("SELECT * FROM analyses WHERE user_id = ? ORDER BY id DESC", (session["user_id"],)).fetchall()
    return render_template("dashboard.html", reports=reports)


@app.get("/reports/<int:report_id>")
@login_required
def saved_report(report_id):
    with get_db() as database:
        report = database.execute("SELECT * FROM analyses WHERE id = ? AND user_id = ?", (report_id, session["user_id"])).fetchone()
    if not report:
        return redirect(url_for("dashboard"))
    return render_template("saved_report.html", report=report, result=json.loads(report["result_json"]))


def owned_report(report_id):
    with get_db() as database:
        return database.execute("SELECT * FROM analyses WHERE id = ? AND user_id = ?", (report_id, session["user_id"])).fetchone()


@app.get("/reports/<int:report_id>/receipt")
@login_required
def report_receipt(report_id):
    report = owned_report(report_id)
    if not report:
        return redirect(url_for("dashboard"))
    return render_template("receipt.html", report=report, result=json.loads(report["result_json"]))


@app.get("/reports/<int:report_id>/export/<string:format_name>")
@login_required
def export_report(report_id, format_name):
    report = owned_report(report_id)
    if not report:
        return jsonify(ok=False, error="Report not found."), 404
    result = json.loads(report["result_json"])
    base_name = f"sourcetrace-report-{report_id}"
    if format_name == "json":
        return Response(json.dumps(result, indent=2), mimetype="application/json",
                        headers={"Content-Disposition": f"attachment; filename={base_name}.json"})
    if format_name == "txt":
        lines = ["SOURCETRACE AI ORIGINALITY RECEIPT", f"Report ID: ST-{report_id:06d}",
                 f"Document: {report['filename']}", f"Words: {report['words']}",
                 f"Similarity: {report['similarity']}%", f"Matched passages: {report['matches']}", ""]
        for index, match in enumerate(result.get("matches", []), 1):
            lines.extend([f"MATCH {index} — {match['score']}%", match["text"],
                          f"Source: {match['source_name']} (page {match['source_page']}, line {match['source_line']})", match["source"], ""])
        return Response("\n".join(lines), mimetype="text/plain",
                        headers={"Content-Disposition": f"attachment; filename={base_name}.txt"})
    if format_name == "csv":
        from io import StringIO
        output = StringIO(); writer = csv.writer(output)
        writer.writerow(["match", "score", "document_page", "document_line", "source", "source_page", "source_line", "document_text", "source_text"])
        for index, match in enumerate(result.get("matches", []), 1):
            writer.writerow([index, match["score"], match["document_page"], match["document_line"], match["source_name"], match["source_page"], match["source_line"], match["text"], match["source"]])
        return Response(output.getvalue(), mimetype="text/csv",
                        headers={"Content-Disposition": f"attachment; filename={base_name}.csv"})
    return jsonify(ok=False, error="Unsupported export format."), 400


@app.get("/reports/<int:report_id>/pdf")
@login_required
def report_pdf(report_id):
    with get_db() as database:
        report = database.execute("SELECT file_path FROM analyses WHERE id = ? AND user_id = ?", (report_id, session["user_id"])).fetchone()
    if not report or not report["file_path"]:
        return jsonify(ok=False, error="Original PDF is not available for this report."), 404
    absolute_path = os.path.abspath(report["file_path"])
    uploads_root = os.path.abspath(UPLOADS_ROOT)
    if os.path.commonpath([absolute_path, uploads_root]) != uploads_root or not os.path.isfile(absolute_path):
        return jsonify(ok=False, error="Original PDF could not be found."), 404
    return send_file(absolute_path, mimetype="application/pdf", download_name="document.pdf")


@app.post("/reports/<int:report_id>/delete")
@login_required
def delete_report(report_id):
    with get_db() as database:
        database.execute("DELETE FROM analyses WHERE id = ? AND user_id = ?", (report_id, session["user_id"]))
    flash("Report deleted.", "info")
    return redirect(url_for("dashboard"))


@app.post("/api/analyze")
def analyze():
    try:
        text = (request.form.get("text") or "").strip()
        document_pages = [text] if text else []
        extracted_document = None
        reference = (request.form.get("reference") or "").strip()
        upload = request.files.get("file")
        if not text and upload and upload.filename:
            extracted_document = extract_upload(upload)
            text = extracted_document["text"].strip()
            document_pages = extracted_document["pages"]
        if not text:
            return jsonify(ok=False, error="Paste text or upload a TXT, PDF, DOCX, or image file."), 400
        words = tokenize(text)
        minimum_words = 5 if extracted_document and extracted_document.get("extension") in IMAGE_EXTENSIONS else MIN_WORDS
        if len(words) < minimum_words:
            if extracted_document and extracted_document.get("extension") in IMAGE_EXTENSIONS:
                return jsonify(ok=False, error="Only a few readable words were detected. Upload a clearer or higher-resolution image."), 400
            return jsonify(ok=False, error=f"Please provide at least {MIN_WORDS} words for a meaningful analysis."), 400
        sources = []
        if reference:
            sources.append({"name": "Pasted reference content", "pages": [reference], "text": reference})
        for reference_file in request.files.getlist("reference_files"):
            if reference_file and reference_file.filename:
                sources.append(extract_upload(reference_file))
        web_discovery = not sources
        if web_discovery:
            sources = discover_web_sources(text)
        combined_reference = "\n".join(source["text"] for source in sources)
        similarity = calculate_similarity(text, combined_reference)
        label, tone = status_for(similarity)
        matches = find_matches(document_pages, sources)
        source_summary = []
        for source in sources:
            source_matches = [match for match in matches if match["source_name"] == source["name"]]
            source_summary.append({"name": source["name"], "domain": source.get("domain", "Uploaded reference"),
                                   "url": source.get("url", ""), "similarity": calculate_similarity(text, source["text"]),
                                   "matches": len(source_matches)})
        source_summary.sort(key=lambda item: item["similarity"], reverse=True)
        writing_signals = analyze_writing_signals(text)
        result = dict(
            ok=True,
            words=len(words),
            characters=len(text),
            similarity=similarity,
            label=label,
            tone=tone,
            matches=matches,
            sources=source_summary,
            document_pages=[{"number": number, "text": page} for number, page in enumerate(document_pages, 1)],
            total_pages=len(document_pages),
            ai_detection={"available": writing_signals["ai_score"] is not None,
                          "score": writing_signals["ai_score"], "label": "Experimental AI-writing signal",
                          "note": writing_signals["note"]},
            grammar={"available": writing_signals["grammar_score"] is not None,
                     "score": writing_signals["grammar_score"], "label": "Grammar & readability"},
            comparison_provided=bool(sources),
            web_discovery=web_discovery,
            note="Similarity score is an automated indicator. Review matched passages and sources before making academic decisions.",
        )
        if session.get("user_id"):
            filename = secure_filename(upload.filename) if upload and upload.filename else "Pasted document"
            with get_db() as database:
                cursor = database.execute("INSERT INTO analyses (user_id, filename, words, pages, similarity, matches, result_json, created_at, ai_score, grammar_score) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
                                          (session["user_id"], filename, len(words), len(document_pages), similarity, len(matches), json.dumps(result), datetime.now(timezone.utc).isoformat(), writing_signals["ai_score"], writing_signals["grammar_score"]))
            result["saved"] = True
            result["report_id"] = cursor.lastrowid
            if extracted_document and extracted_document.get("extension") == "pdf":
                user_upload_dir = os.path.join(UPLOADS_ROOT, str(session["user_id"]))
                os.makedirs(user_upload_dir, exist_ok=True)
                stored_path = os.path.join(user_upload_dir, f"report-{cursor.lastrowid}.pdf")
                with open(stored_path, "wb") as stored_pdf:
                    stored_pdf.write(extracted_document["data"])
                with get_db() as database:
                    database.execute("UPDATE analyses SET file_path = ? WHERE id = ? AND user_id = ?", (stored_path, cursor.lastrowid, session["user_id"]))
                result["pdf_available"] = True
        else:
            result["saved"] = False
        return jsonify(result)
    except ValueError as error:
        return jsonify(ok=False, error=str(error)), 400


@app.errorhandler(RequestEntityTooLarge)
def file_too_large(_error):
    return jsonify(ok=False, error="The file is too large. Maximum upload size is 10 MB."), 413


if __name__ == "__main__":
    init_db()
    app.run(debug=True)


init_db()
